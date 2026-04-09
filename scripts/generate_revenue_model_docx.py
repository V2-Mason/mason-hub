"""Generate Growth Memo Revenue Model decision document as Word."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from md_to_docx_batch import add_markdown_to_doc

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04/cross-comparison/growth_memo_revenue_model.md")
OUTPUT = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04/cross-comparison/Growth_Memo_Revenue_Model.docx")


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover title
    title = doc.add_heading("Growth Memo 变现模式决策", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("方案 D: 80% 粉丝经济 + 20% 真工具党评测")
    sr.font.size = Pt(13)
    sr.italic = True
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("决策日期: 2026-04-08  |  版本: v1.0  |  状态: Mason 已确认")
    mr.italic = True
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Status box
    status_heading = doc.add_paragraph()
    sh = status_heading.add_run("本文档的作用")
    sh.bold = True
    sh.font.size = Pt(12)

    for line in [
        "回答 Mason 的核心战略问题: '资金来源应该是粉丝还是金主?'",
        "与 Strategic_Impact_Report.docx 和 bilibili-mason-target.md preset 并列, 作为路线依据",
        "6-12 个月后根据数据重新评估, 决定是否切换到方案 C (双账号)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.4)
        p.add_run(line).font.size = Pt(10)

    doc.add_page_break()

    # Body from markdown
    md_text = SOURCE.read_text(encoding="utf-8")

    # Skip the first H1
    lines = md_text.split("\n")
    filtered_lines = []
    skip_first_h1 = True
    for line in lines:
        if skip_first_h1 and line.startswith("# "):
            skip_first_h1 = False
            continue
        filtered_lines.append(line)
    md_text_body = "\n".join(filtered_lines)

    add_markdown_to_doc(doc, md_text_body, base_heading_level=1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    size_kb = OUTPUT.stat().st_size / 1024
    print(f"[OK] Saved: {OUTPUT}")
    print(f"[OK] Size: {size_kb:.1f} KB")
    print(f"[OK] Paragraphs: {len(doc.paragraphs)}")
    print(f"[OK] Tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
