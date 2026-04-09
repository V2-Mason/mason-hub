"""
Generate the Mine Evidence Plan as a Word document for Mason to review.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from md_to_docx_batch import add_markdown_to_doc

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04/cross-comparison/mine_evidence_plan.md")
OUTPUT = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04/cross-comparison/Mine_Evidence_Plan.docx")


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover title
    title = doc.add_heading("评论挖掘执行计划", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("5 路径参数确认清单 — 等待 Mason 审批")
    subtitle_run.font.size = Pt(13)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("生成日期: 2026-04-08 | 版本: v0.1")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Read source markdown
    md_text = SOURCE.read_text(encoding="utf-8")

    # Skip the first H1 since we added the cover
    lines = md_text.split("\n")
    filtered_lines = []
    skip_first_h1 = True
    for line in lines:
        if skip_first_h1 and line.startswith("# "):
            skip_first_h1 = False
            continue
        filtered_lines.append(line)
    md_text_body = "\n".join(filtered_lines)

    add_markdown_to_doc(doc, md_text_body, base_heading_level=1)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    size_kb = OUTPUT.stat().st_size / 1024
    print(f"[OK] Saved: {OUTPUT}")
    print(f"[OK] Size: {size_kb:.1f} KB")
    print(f"[OK] Paragraphs: {len(doc.paragraphs)}")
    print(f"[OK] Tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
