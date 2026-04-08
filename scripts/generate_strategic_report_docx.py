"""
Generate the Strategic Impact Report as a single Word document.
Uses the markdown parser from md_to_docx_batch.py and applies custom styling.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from md_to_docx_batch import add_markdown_to_doc

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04/cross-comparison/strategic_impact_report.md")
OUTPUT = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04/cross-comparison/Strategic_Impact_Report.docx")


def build():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    # Set page margins (slightly narrower for more content)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover title
    title = doc.add_heading("Growth Memo 战略影响报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("三层数据汇总与决策建议")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("生成日期: 2026-04-08 | 版本: v0.1")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Data basis box
    doc.add_paragraph()
    basis_heading = doc.add_paragraph()
    basis_run = basis_heading.add_run("数据基础")
    basis_run.bold = True
    basis_run.font.size = Pt(12)

    basis_items = [
        "AUDIENCE_PERSONAS: 272 条 B 站评论 + 5 人群画像",
        "官方行业报告: 克劳锐 / 新榜 / 千瓜 / 益普索 / 36氪 / 澎湃 (2024-2025)",
        "赛道侦察报告: 16 关键词 / 301 视频 / 272 评论 / 8 需求板块",
        "今日 4 账号 Grid: 4 × 57 = 228 视频 × 9 check = 2052 格分析",
    ]
    for item in basis_items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.4)
        p.add_run(item).font.size = Pt(10)

    doc.add_paragraph()

    # Document purpose
    purpose_heading = doc.add_paragraph()
    purpose_run = purpose_heading.add_run("文档目的 — 回答两个元问题")
    purpose_run.bold = True
    purpose_run.font.size = Pt(12)

    q1 = doc.add_paragraph(style="List Number")
    q1.paragraph_format.left_indent = Inches(0.4)
    q1.add_run("AUDIENCE_PERSONAS 是怎么分析出来的? 数据基础是什么?").font.size = Pt(10.5)

    q2 = doc.add_paragraph(style="List Number")
    q2.paragraph_format.left_indent = Inches(0.4)
    q2.add_run("之前的行业/赛道报告 + 今日 4 账号 Grid 合起来, 对 Growth Memo 路线到底意味着什么?").font.size = Pt(10.5)

    doc.add_page_break()

    # Read source markdown
    md_text = SOURCE.read_text(encoding="utf-8")

    # Skip the duplicated title (first H1) since we already added a cover
    lines = md_text.split("\n")
    filtered_lines = []
    skip_first_h1 = True
    for line in lines:
        if skip_first_h1 and line.startswith("# "):
            skip_first_h1 = False
            continue
        filtered_lines.append(line)
    md_text_body = "\n".join(filtered_lines)

    # Also skip the initial meta block (generation date, data basis lines) since we've shown those
    # Split at first '---' and keep everything after
    if "\n---\n" in md_text_body:
        parts = md_text_body.split("\n---\n", 1)
        if len(parts) == 2:
            md_text_body = parts[1]

    # Add the body content
    add_markdown_to_doc(doc, md_text_body, base_heading_level=1)

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    # Stats
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"[OK] Saved: {OUTPUT}")
    print(f"[OK] Size: {size_kb:.1f} KB")
    print(f"[OK] Paragraphs: {len(doc.paragraphs)}")
    print(f"[OK] Tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
