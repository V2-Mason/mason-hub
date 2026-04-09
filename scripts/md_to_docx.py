#!/usr/bin/env python3
"""
Minimal markdown -> docx converter, tuned for Mason's growth-memo analysis files.
Handles: headings (#-####), bold **, italic *, lists (-, *, N.), tables (|), blockquote (>),
horizontal rule (---), code spans (`), paragraphs, and hyperlinks [text](url) (rendered as text only).

Usage:
  python scripts/md_to_docx.py <input.md> [output.docx]

If output is omitted, uses <input_stem>.docx next to the input.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Microsoft YaHei"

HEADING_SIZES = {1: 22, 2: 18, 3: 15, 4: 13}
BODY_SIZE = 11
SMALL_SIZE = 10

HEADING_COLORS = {
    1: RGBColor(0x1F, 0x4E, 0x79),
    2: RGBColor(0x2E, 0x75, 0xB6),
    3: RGBColor(0x2E, 0x75, 0xB6),
    4: RGBColor(0x50, 0x50, 0x50),
}


def _set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # East Asian font (for Chinese glyphs)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)


# ----- Inline parser -----
# Tokens produced: list of (text, bold, italic, code)
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))"
)


def parse_inline(text):
    """Split a string into styled segments. Returns list of dicts."""
    out = []
    idx = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > idx:
            out.append({"text": text[idx:m.start()], "bold": False, "italic": False, "code": False})
        tok = m.group(0)
        if tok.startswith("**") and tok.endswith("**"):
            out.append({"text": tok[2:-2], "bold": True, "italic": False, "code": False})
        elif tok.startswith("`") and tok.endswith("`"):
            out.append({"text": tok[1:-1], "bold": False, "italic": False, "code": True})
        elif tok.startswith("[") and "](" in tok:
            # [text](url) -> render as text only (underlined would need hyperlink)
            inner = tok[1:-1]
            t, _, u = inner.partition("](")
            out.append({"text": t, "bold": False, "italic": True, "code": False})
        idx = m.end()
    if idx < len(text):
        out.append({"text": text[idx:], "bold": False, "italic": False, "code": False})
    return out


def add_inline_runs(paragraph, text, base_size=BODY_SIZE, base_bold=False):
    segs = parse_inline(text)
    if not segs:
        run = paragraph.add_run("")
        _set_run_font(run, size=base_size, bold=base_bold)
        return
    for seg in segs:
        run = paragraph.add_run(seg["text"])
        _set_run_font(
            run,
            size=base_size,
            bold=base_bold or seg["bold"],
            italic=seg["italic"],
            color=RGBColor(0xC0, 0x00, 0x00) if seg["code"] else None,
        )


# ----- Block parser -----
def parse_blocks(lines):
    """Yield block tuples: (kind, payload)."""
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            yield ("hr", None)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            yield ("heading", (level, m.group(2).strip()))
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            j = i + 1
            code_lines = []
            while j < n and not lines[j].strip().startswith("```"):
                code_lines.append(lines[j])
                j += 1
            yield ("code", (lang, code_lines))
            i = j + 1
            continue

        # Table (line containing |, followed by separator row)
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?[\s:\-]*\|[\s:\-|]*$", lines[i + 1]):
            header = _split_table_row(lines[i])
            i += 2  # skip header + separator
            rows = []
            while i < n:
                row_line = lines[i]
                if "|" not in row_line.strip() or not row_line.strip():
                    break
                rows.append(_split_table_row(row_line))
                i += 1
            yield ("table", (header, rows))
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            yield ("quote", " ".join(quote_lines))
            continue

        # Unordered list
        if re.match(r"^[\-\*]\s+", stripped):
            items = []
            while i < n and re.match(r"^[\-\*]\s+", lines[i].strip()):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            yield ("ulist", items)
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            yield ("olist", items)
            continue

        # Paragraph: gather until blank or new block marker
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r"^#{1,6}\s", nxt):
                break
            if re.match(r"^-{3,}$", nxt):
                break
            if nxt.startswith("```"):
                break
            if "|" in nxt and i + 1 < n and re.match(r"^\s*\|?[\s:\-]*\|[\s:\-|]*$", lines[i + 1]):
                break
            if re.match(r"^[\-\*]\s+", nxt):
                break
            if re.match(r"^\d+\.\s+", nxt):
                break
            if nxt.startswith(">"):
                break
            para_lines.append(nxt)
            i += 1
        yield ("para", " ".join(para_lines))


def _split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


# ----- DOCX writer -----
def write_docx(md_path: Path, out_path: Path, title: str | None = None):
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)

    # Page margins (narrower so tables fit)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    for kind, payload in parse_blocks(lines):
        if kind == "heading":
            level, text = payload
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            p.paragraph_format.space_after = Pt(4)
            size = HEADING_SIZES.get(level, BODY_SIZE)
            color = HEADING_COLORS.get(level)
            # Parse inline in headings too (for bold emphasis)
            segs = parse_inline(text)
            if not segs:
                run = p.add_run(text)
                _set_run_font(run, size=size, bold=True, color=color)
            else:
                for seg in segs:
                    run = p.add_run(seg["text"])
                    _set_run_font(run, size=size, bold=True, italic=seg["italic"], color=color)

        elif kind == "para":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, payload, base_size=BODY_SIZE)

        elif kind == "ulist":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, item, base_size=BODY_SIZE)

        elif kind == "olist":
            for item in payload:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, item, base_size=BODY_SIZE)

        elif kind == "table":
            header, rows = payload
            if not header:
                continue
            ncols = len(header)
            tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            tbl.autofit = True
            # Header
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(header[:ncols]):
                cell_para = hdr_cells[i].paragraphs[0]
                cell_para.paragraph_format.space_after = Pt(0)
                segs = parse_inline(h)
                if not segs:
                    run = cell_para.add_run(h)
                    _set_run_font(run, size=SMALL_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                else:
                    for seg in segs:
                        run = cell_para.add_run(seg["text"])
                        _set_run_font(run, size=SMALL_SIZE, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                # Shade header cell
                tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2E75B6")
                tcPr.append(shd)
            # Body
            for r_idx, row in enumerate(rows):
                cells = tbl.rows[r_idx + 1].cells
                for c_idx in range(ncols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    cell_para = cells[c_idx].paragraphs[0]
                    cell_para.paragraph_format.space_after = Pt(0)
                    add_inline_runs(cell_para, val, base_size=SMALL_SIZE)
            # Spacer
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(payload)
            _set_run_font(run, size=BODY_SIZE, italic=True, color=RGBColor(0x66, 0x66, 0x66))

        elif kind == "code":
            lang, code_lines = payload
            for code_line in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run(code_line)
                _set_run_font(run, size=SMALL_SIZE, color=RGBColor(0x30, 0x30, 0x30))
                run.font.name = "Consolas"

        elif kind == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 60)
            _set_run_font(run, size=BODY_SIZE, color=RGBColor(0xCC, 0xCC, 0xCC))

    doc.save(str(out_path))
    print(f"Wrote: {out_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Input not found: {md_path}")
        sys.exit(1)
    if len(sys.argv) >= 3:
        out_path = Path(sys.argv[2])
    else:
        out_path = md_path.with_suffix(".docx")
    write_docx(md_path, out_path)


if __name__ == "__main__":
    main()
