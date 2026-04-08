"""
Batch convert analysis markdown files to a single Word document.
Targets the 4 Tier 1 Bilibili creator grid analysis outputs.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path("c:/Users/hangn/projects/mason-hub/accounts/growth-memo/content/test-001/assets/reference/batch-2026-04")

# Ordered list: (section_title, list of (subtitle, filepath))
SECTIONS = [
    ("Section 1 — 程序员鱼皮 (Tier 1)", [
        ("1.1 Grid (57 × 9)", BASE / "tier1-programmer-yupi/analysis/grid.md"),
        ("1.2 Hit Rate (Phase 4)", BASE / "tier1-programmer-yupi/analysis/hit_rate.md"),
        ("1.3 Patterns (Phase 5)", BASE / "tier1-programmer-yupi/analysis/patterns.md"),
    ]),
    ("Section 2 — ezindie (Tier 1)", [
        ("2.1 Grid (57 × 9)", BASE / "tier1-ezindie/analysis/grid.md"),
        ("2.2 Hit Rate (Phase 4)", BASE / "tier1-ezindie/analysis/hit_rate.md"),
        ("2.3 Patterns (Phase 5)", BASE / "tier1-ezindie/analysis/patterns.md"),
    ]),
    ("Section 3 — 小Lin说 (Tier 2)", [
        ("3.1 Scan A Creator Preset — Grid", BASE / "tier2-xiaolin-shuo/analysis/scan_a_creator_preset/grid.md"),
        ("3.2 Scan A Creator Preset — Hit Rate", BASE / "tier2-xiaolin-shuo/analysis/scan_a_creator_preset/hit_rate.md"),
        ("3.3 Scan A Creator Preset — Patterns", BASE / "tier2-xiaolin-shuo/analysis/scan_a_creator_preset/patterns.md"),
        ("3.4 Scan B Dev Preset — Grid", BASE / "tier2-xiaolin-shuo/analysis/scan_b_dev_preset/grid.md"),
        ("3.5 Scan B Dev Preset — Hit Rate", BASE / "tier2-xiaolin-shuo/analysis/scan_b_dev_preset/hit_rate.md"),
        ("3.6 Scan B Dev Preset — Patterns", BASE / "tier2-xiaolin-shuo/analysis/scan_b_dev_preset/patterns.md"),
        ("3.7 Track Verdict (A vs B 对比)", BASE / "tier2-xiaolin-shuo/analysis/track_verdict.md"),
    ]),
    ("Section 4 — 巫师财经 (Tier 2)", [
        ("4.1 Scan A Creator Preset — Grid", BASE / "tier2-wushi-finance/analysis/scan_a_creator_preset/grid.md"),
        ("4.2 Scan A Creator Preset — Hit Rate", BASE / "tier2-wushi-finance/analysis/scan_a_creator_preset/hit_rate.md"),
        ("4.3 Scan A Creator Preset — Patterns", BASE / "tier2-wushi-finance/analysis/scan_a_creator_preset/patterns.md"),
        ("4.4 Scan B Dev Preset — Grid", BASE / "tier2-wushi-finance/analysis/scan_b_dev_preset/grid.md"),
        ("4.5 Scan B Dev Preset — Hit Rate", BASE / "tier2-wushi-finance/analysis/scan_b_dev_preset/hit_rate.md"),
        ("4.6 Scan B Dev Preset — Patterns", BASE / "tier2-wushi-finance/analysis/scan_b_dev_preset/patterns.md"),
        ("4.7 Track Verdict (A vs B 对比)", BASE / "tier2-wushi-finance/analysis/track_verdict.md"),
    ]),
]

OUTPUT = BASE / "cross-comparison/4tier1_analysis_all.docx"

# --- Markdown parsing ---

def parse_inline(text):
    """Return list of (text, is_bold, is_italic, is_code) tuples."""
    tokens = []
    i = 0
    n = len(text)
    while i < n:
        # Bold: **...** or __...__
        if text[i:i+2] == "**":
            end = text.find("**", i + 2)
            if end != -1:
                tokens.append((text[i+2:end], True, False, False))
                i = end + 2
                continue
        # Italic: *...* (single) or _..._ (but skip if preceded/followed by non-space word char for _)
        if text[i] == "*" and (i + 1 < n and text[i+1] != "*"):
            end = text.find("*", i + 1)
            if end != -1 and end > i + 1:
                tokens.append((text[i+1:end], False, True, False))
                i = end + 1
                continue
        # Code: `...`
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end != -1:
                tokens.append((text[i+1:end], False, False, True))
                i = end + 1
                continue
        # Plain char accumulation
        j = i
        while j < n and text[j] not in "*`":
            j += 1
        if j > i:
            tokens.append((text[i:j], False, False, False))
            i = j
        else:
            tokens.append((text[i], False, False, False))
            i += 1
    # Merge adjacent plain tokens
    merged = []
    for tok in tokens:
        if merged and merged[-1][1] == tok[1] and merged[-1][2] == tok[2] and merged[-1][3] == tok[3]:
            merged[-1] = (merged[-1][0] + tok[0], tok[1], tok[2], tok[3])
        else:
            merged.append(tok)
    return merged


def add_formatted_text(paragraph, text):
    for token, bold, italic, code in parse_inline(text):
        run = paragraph.add_run(token)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def parse_table_row(line):
    # Split on | but ignore leading/trailing |
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line):
    line = line.strip()
    if not line.startswith("|"):
        return False
    cells = parse_table_row(line)
    return all(re.fullmatch(r":?-+:?", c) for c in cells)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_markdown_to_doc(doc, md_text, base_heading_level=2):
    """Parse markdown and append to doc. base_heading_level shifts md H1 to that level."""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    in_code = False
    code_buffer = []

    while i < n:
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                # Close
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Table detection: current line has | and next line is separator
        if line.strip().startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            headers = parse_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|") and not is_table_separator(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            # Build docx table
            ncols = len(headers)
            tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            tbl.autofit = True
            hdr_cells = tbl.rows[0].cells
            for ci, h in enumerate(headers):
                cell = hdr_cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(8)
                set_cell_shading(cell, "D9E2F3")
            for ri, row in enumerate(rows):
                row_cells = tbl.rows[ri + 1].cells
                for ci in range(ncols):
                    txt = row[ci] if ci < len(row) else ""
                    cell = row_cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    # Highlight PASS green / FAIL red in cell text
                    if "PASS" in txt and "FAIL" not in txt:
                        run = p.add_run(txt)
                        run.font.color.rgb = RGBColor(0x0B, 0x6A, 0x0B)
                        run.font.size = Pt(7)
                    elif txt.strip().startswith("FAIL") or txt.strip() == "FAIL":
                        run = p.add_run(txt)
                        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
                        run.font.size = Pt(7)
                    else:
                        add_formatted_text(p, txt)
                        for r in p.runs:
                            if r.font.size is None:
                                r.font.size = Pt(7)
            doc.add_paragraph()  # spacing after table
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading_level = min(base_heading_level + level - 1, 9)
            try:
                h = doc.add_heading("", level=heading_level)
                add_formatted_text(h, text)
            except Exception:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$|^\*{3,}$", line.strip()):
            p = doc.add_paragraph()
            p.add_run("_" * 60).italic = True
            i += 1
            continue

        # List item (unordered)
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent)
            add_formatted_text(p, text)
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[1:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1


def build_document():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Batch Recon 2026-04 Tier 1 四账号 Grid 分析合集", level=0)
    for run in title.runs:
        run.font.name = "Microsoft YaHei"

    meta = doc.add_paragraph()
    meta.add_run("生成日期: 2026-04-08 | 账号: 程序员鱼皮 / ezindie / 小Lin说 / 巫师财经 | 总样本: 4 × 57 = 228 条视频").italic = True

    doc.add_paragraph(
        "本文档合并了 Batch Recon 2026-04 session 中对 4 个 B 站 Tier 1 头部账号的 Grid 分析产物。"
        "鱼皮和 ezindie 用 bilibili-dev-content 单 preset 扫描，小Lin说和巫师财经用 creator-preset (Scan A) + dev-preset (Scan B) 双扫对比。"
        "每个账号均包含 Phase 2-3 的 57×9 Grid、Phase 4 的命中率统计、Phase 5 的规律提炼。双扫账号额外附 track_verdict 对比裁定。"
    )

    # Table of contents placeholder
    doc.add_heading("目录", level=1)
    for sec_title, items in SECTIONS:
        p = doc.add_paragraph()
        run = p.add_run(sec_title)
        run.bold = True
        for sub_title, _ in items:
            sub_p = doc.add_paragraph(style="List Bullet")
            sub_p.paragraph_format.left_indent = Inches(0.5)
            sub_p.add_run(sub_title)

    doc.add_page_break()

    # Sections
    for sec_title, items in SECTIONS:
        doc.add_heading(sec_title, level=1)
        for sub_title, path in items:
            doc.add_heading(sub_title, level=2)
            if not path.exists():
                p = doc.add_paragraph()
                run = p.add_run(f"[文件缺失: {path}]")
                run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
                continue
            md_text = path.read_text(encoding="utf-8")
            add_markdown_to_doc(doc, md_text, base_heading_level=3)
            doc.add_paragraph()
        doc.add_page_break()

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"[OK] Saved: {OUTPUT}")
    print(f"[OK] Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
